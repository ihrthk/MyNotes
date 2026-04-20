#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
还款计划表Word文档生成器
创建详细的还款计划表
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime, timedelta
import os

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tcBorders = OxmlElement('w:tcBorders')
    for edge in ['top', 'left', 'bottom', 'right']:
        if edge in kwargs:
            edge_element = OxmlElement(f'w:{edge}')
            edge_element.set(qn('w:val'), kwargs[edge].get('val', 'single'))
            edge_element.set(qn('w:sz'), str(kwargs[edge].get('sz', 4)))
            edge_element.set(qn('w:color'), kwargs[edge].get('color', '000000'))
            tcBorders.append(edge_element)

    tcPr.ap pend(tcBorders)

def set_chinese_font(run):
    """设置中文字体"""
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def calculate_monthly_rate(principal, monthly_payment, total_months):
    """
    使用牛顿迭代法求解实际月利率（更精确）
    """
    r = 0.006  # 初始猜测值 0.6%

    for iteration in range(1000):
        # 等额本息公式：月供 = [本金 × r × (1+r)^n] / [(1+r)^n - 1]
        A = (1 + r) ** total_months
        B = A - 1

        numerator = principal * r * A
        denominator = B
        calculated_payment = numerator / denominator

        # 计算误差
        error = calculated_payment - monthly_payment

        # 如果误差足够小，返回结果
        if abs(error) < 0.000001:
            return r

        # 计算导数（用于牛顿法）
        C = r * total_months * (1 + r) ** (total_months - 1)
        derivative = principal * (B + C) / (B ** 2)

        # 牛顿迭代
        r = r - error / derivative

    return r

def calculate_repayment_schedule(principal, annual_rate, total_months, monthly_payment):
    """
    计算还款计划表（使用固定月供）

    参数:
    - principal: 本金
    - annual_rate: 年化利率（小数形式，如0.0796，仅用于显示）
    - total_months: 总期数
    - monthly_payment: 每月还款额（固定值）

    返回: 还款计划列表和实际年利率
    """
    schedule = []
    remaining_principal = principal

    # 计算实际月利率（基于固定月供3000元）
    actual_monthly_rate = calculate_monthly_rate(principal, monthly_payment, total_months)
    actual_annual_rate = actual_monthly_rate * 12

    print(f"固定月供: ¥{monthly_payment:.2f} 元")
    print(f"实际月利率: {actual_monthly_rate*100:.6f}%")
    print(f"实际年利率: {actual_annual_rate*100:.3f}%")

    for month in range(1, total_months + 1):
        # 计算当期利息（使用实际月利率）
        interest = remaining_principal * actual_monthly_rate

        # 计算当期本金
        # 最后一期调整
        if month == total_months:
            principal_payment = remaining_principal
            total_payment = principal_payment + interest
        else:
            principal_payment = monthly_payment - interest
            total_payment = monthly_payment

        # 更新剩余本金
        remaining_principal -= principal_payment
        if remaining_principal < 0.01:  # 避免浮点数误差
            remaining_principal = 0

        # 计算还款日期
        year = 2026 + (month - 1) // 12
        month_num = (month - 1) % 12 + 2
        if month_num > 12:
            month_num -= 12
            year += 1

        schedule.append({
            '期数': month,
            '还款日期': f"{year}年{month_num}月15日",
            '还款额': round(total_payment, 2),
            '偿还本金': round(principal_payment, 2),
            '偿还利息': round(interest, 2),
            '剩余本金': round(remaining_principal, 2)
        })

    return schedule, actual_annual_rate

def create_repayment_schedule_docx(filename="还款计划表.docx"):
    """创建还款计划表Word文档"""

    # 创建文档对象
    doc = Document()

    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        # 设置为横向
        section.page_height = Inches(8.27)
        section.page_width = Inches(11.69)

    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.add_run('还款计划表')
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    set_chinese_font(title_run)

    # 计算还款计划（需要先计算，获取实际年利率）
    principal = 250000
    annual_rate = 0.0796
    total_months = 120
    monthly_payment = 3000

    schedule, actual_annual_rate = calculate_repayment_schedule(principal, annual_rate, total_months, monthly_payment)

    # 借款基本信息
    doc.add_paragraph()

    info_table = doc.add_table(rows=5, cols=4)
    info_table.alignment = 1

    # 设置表格样式
    for row in info_table.rows:
        for cell in row.cells:
            set_cell_border(cell, top={'val': 'single'}, bottom={'val': 'single'},
                          left={'val': 'single'}, right={'val': 'single'})

    # 填充信息
    info_data = [
        ['借款人', '李广', '出借人', '张利顺'],
        ['借款本金', '¥250,000.00', '年化利率', f'{actual_annual_rate*100:.3f}%'],
        ['还款期数', '120期', '还款方式', '等额本息'],
        ['首期还款日', '2026年2月15日', '还款日', '每月15日'],
        ['每月还款额', '¥3,000.00', '还款总额', '¥360,000.00']
    ]

    for i, row_data in enumerate(info_data):
        for j, cell_data in enumerate(row_data):
            cell = info_table.rows[i].cells[j]
            cell.text = str(cell_data)
            # 设置单元格对齐
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
                set_chinese_font(run)
                if j % 2 == 0:  # 标签列加粗
                    run.font.bold = True

    doc.add_paragraph()

    # 添加说明
    note = doc.add_paragraph()
    note_run = note.add_run('注：本还款计划表为借款合同的组成部分，与《借条》具有同等法律效力。')
    note_run.font.size = Pt(9)
    note_run.font.color.rgb = RGBColor(128, 128, 128)
    set_chinese_font(note_run)

    doc.add_paragraph()

    # 创建还款计划表（分两页显示）
    # 第一页：1-60期
    doc.add_paragraph()
    page1_title = doc.add_paragraph()
    page1_title_run = page1_title.add_run('还款明细（第1-60期）')
    page1_title_run.font.size = Pt(12)
    page1_title_run.font.bold = True
    set_chinese_font(page1_title_run)

    doc.add_paragraph()

    # 创建表格
    table = doc.add_table(rows=1, cols=6)
    table.alignment = 1

    # 设置表头
    headers = ['期数', '还款日期', '还款额', '偿还本金', '偿还利息', '剩余本金']
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        cell = header_cells[i]
        cell.text = header
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(9)
            run.font.bold = True
            set_chinese_font(run)
        # 设置表头背景色
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'D9D9D9')
        cell._element.get_or_add_tcPr().append(shading_elm)

    # 添加数据行
    for item in schedule[:60]:
        row_cells = table.add_row().cells
        row_data = [
            str(item['期数']),
            item['还款日期'],
            f"¥{item['还款额']:,.2f}",
            f"¥{item['偿还本金']:,.2f}",
            f"¥{item['偿还利息']:,.2f}",
            f"¥{item['剩余本金']:,.2f}"
        ]
        for i, data in enumerate(row_data):
            cell = row_cells[i]
            cell.text = data
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(8)
                set_chinese_font(run)

    # 设置所有单元格边框
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top={'val': 'single'}, bottom={'val': 'single'},
                          left={'val': 'single'}, right={'val': 'single'})

    # 添加分页符
    doc.add_page_break()

    # 第二页标题
    page2_title = doc.add_paragraph()
    page2_title_run = page2_title.add_run('还款明细（第61-120期）')
    page2_title_run.font.size = Pt(12)
    page2_title_run.font.bold = True
    set_chinese_font(page2_title_run)

    doc.add_paragraph()

    # 创建第二页表格
    table2 = doc.add_table(rows=1, cols=6)
    table2.alignment = 1

    # 设置表头
    header_cells2 = table2.rows[0].cells
    for i, header in enumerate(headers):
        cell = header_cells2[i]
        cell.text = header
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(9)
            run.font.bold = True
            set_chinese_font(run)
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'D9D9D9')
        cell._element.get_or_add_tcPr().append(shading_elm)

    # 添加数据行
    for item in schedule[60:]:
        row_cells = table2.add_row().cells
        row_data = [
            str(item['期数']),
            item['还款日期'],
            f"¥{item['还款额']:,.2f}",
            f"¥{item['偿还本金']:,.2f}",
            f"¥{item['偿还利息']:,.2f}",
            f"¥{item['剩余本金']:,.2f}"
        ]
        for i, data in enumerate(row_data):
            cell = row_cells[i]
            cell.text = data
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(8)
                set_chinese_font(run)

    # 设置所有单元格边框
    for row in table2.rows:
        for cell in row.cells:
            set_cell_border(cell, top={'val': 'single'}, bottom={'val': 'single'},
                          left={'val': 'single'}, right={'val': 'single'})

    # 保存文档
    doc.save(filename)
    print(f"✓ 还款计划表Word文档已生成: {filename}")
    print(f"✓ 文件位置: {os.path.abspath(filename)}")
    print(f"✓ 共{total_months}期，分两页显示")

if __name__ == "__main__":
    print("=" * 60)
    print("还款计划表Word文档生成器")
    print("=" * 60)

    # 检查是否安装了python-docx
    try:
        import docx
    except ImportError:
        print("❌ 错误: 未安装python-docx库")
        print("请运行: pip3 install python-docx")
        exit(1)

    # 生成还款计划表
    create_repayment_schedule_docx()
    print("\n✓ 完成！")
