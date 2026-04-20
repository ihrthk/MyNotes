#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
借条Word文档生成器
创建专业的借条Word文档
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
import os

def create_iou_docx(filename="借条.docx"):
    """创建借条Word文档"""

    # 创建文档对象
    doc = Document()

    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.add_run('借 条')
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.name = '黑体'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    # 前言
    intro = doc.add_paragraph()
    intro_run = intro.add_run('因资金周转需要，今本人 ')
    intro_run.font.size = Pt(12)
    set_chinese_font(intro_run)

    intro_run_borrower = intro.add_run('李广')
    intro_run_borrower.font.bold = True
    intro_run_borrower.font.size = Pt(12)
    set_chinese_font(intro_run_borrower)

    intro_run2 = intro.add_run('（借款人）收到 ')
    intro_run2.font.size = Pt(12)
    set_chinese_font(intro_run2)

    intro_run_lender = intro.add_run('张利顺')
    intro_run_lender.font.bold = True
    intro_run_lender.font.size = Pt(12)
    set_chinese_font(intro_run_lender)

    intro_run3 = intro.add_run('（出借人）通过银行转账方式出借的款项共计人民币')
    intro_run3.font.size = Pt(12)
    set_chinese_font(intro_run3)

    intro_run_amount = intro.add_run('贰拾伍万元整（¥250,000.00）')
    intro_run_amount.font.bold = True
    intro_run_amount.font.size = Pt(12)
    set_chinese_font(intro_run_amount)

    intro_run4 = intro.add_run('。具体支付情况如下，本人确认已全部收到：')
    intro_run4.font.size = Pt(12)
    set_chinese_font(intro_run4)

    # 转账明细
    payment_details = [
        '2022年6月8日，收到出借人通过其本人名下工商银行账户（卡号：6212260200147622993）转账支付的伍万元整（¥50,000.00）；',
        '2022年6月9日，收到出借人通过其本人名下上述同一账户转账支付的伍万元整（¥50,000.00）；',
        '2022年6月10日，收到出借人通过其本人名下上述同一账户转账支付的伍万元整（¥50,000.00）；',
        '2024年3月8日，收到出借人通过其本人名下上述同一账户转账支付的壹拾万元整（¥100,000.00）。'
    ]

    for detail in payment_details:
        p = doc.add_paragraph(detail, style='List Number')
        p.paragraph_format.left_indent = Inches(0.5)
        for run in p.runs:
            run.font.size = Pt(11)
            set_chinese_font(run)

    # 总额确认
    total = doc.add_paragraph()
    total_run = total.add_run('双方确认，上述四笔转账共同构成一笔总额为人民币')
    total_run.font.size = Pt(12)
    total_run.font.bold = True
    set_chinese_font(total_run)

    total_run2 = total.add_run('贰拾伍万元整（¥250,000.00）')
    total_run2.font.size = Pt(12)
    total_run2.font.bold = True
    set_chinese_font(total_run2)

    total_run3 = total.add_run('的借款。')
    total_run3.font.size = Pt(12)
    total_run3.font.bold = True
    set_chinese_font(total_run3)

    # 还款条款
    repayment = doc.add_paragraph()
    repayment_text = (
        '还款方式为等额本息，分120期按月偿还，每月还款额为人民币叁仟元整（¥3,000.00），该金额为固定月供。'
        '根据上述本金、期限及月供金额折算，本借款的实际年化利率约为7.748%。'
        '首期还款日为2026年2月15日，此后每月15日为还款日。'
        '具体每期应还本金、利息及剩余本金详见附件《还款计划表》，该附件与本借条具有同等法律效力。'
    )
    repayment_run = repayment.add_run(repayment_text)
    repayment_run.font.size = Pt(12)
    set_chinese_font(repayment_run)

    # 提前还款
    early_payment_title = doc.add_paragraph()
    early_payment_title_run = early_payment_title.add_run('第一条  提前还款')
    early_payment_title_run.font.size = Pt(12)
    early_payment_title_run.font.bold = True
    set_chinese_font(early_payment_title_run)

    early_payment = doc.add_paragraph()
    early_payment_run = early_payment.add_run(
        '借款人有权提前清偿全部剩余借款。若提前还款，借款人只需清偿提前还款日当日的剩余本金，无需支付该日之后的任何利息。'
    )
    early_payment_run.font.size = Pt(12)
    set_chinese_font(early_payment_run)

    # 违约责任
    default_title = doc.add_paragraph()
    default_title_run = default_title.add_run('第二条  违约责任')
    default_title_run.font.size = Pt(12)
    default_title_run.font.bold = True
    set_chinese_font(default_title_run)

    default = doc.add_paragraph()
    default_run = default.add_run(
        '若借款人未按附件《还款计划表》约定的日期足额偿还任何一期款项，则全部借款立即提前到期，'
        '出借人有权要求借款人一次性清偿剩余全部本金。'
    )
    default_run.font.size = Pt(12)
    set_chinese_font(default_run)

    # 争议解决
    dispute_title = doc.add_paragraph()
    dispute_title_run = dispute_title.add_run('第三条  争议解决')
    dispute_title_run.font.size = Pt(12)
    dispute_title_run.font.bold = True
    set_chinese_font(dispute_title_run)

    dispute = doc.add_paragraph()
    dispute_run = dispute.add_run(
        '因履行本借条所产生的任何争议，双方应协商解决；协商不成的，任何一方均有权向出借人住所地人民法院提起诉讼。'
    )
    dispute_run.font.size = Pt(12)
    set_chinese_font(dispute_run)

    # 当事人信息
    doc.add_paragraph()  # 空行

    # 借款人信息
    borrower_info = doc.add_paragraph()
    borrower_info_run = borrower_info.add_run('借款人信息：')
    borrower_info_run.font.size = Pt(12)
    borrower_info_run.font.bold = True
    set_chinese_font(borrower_info_run)

    borrower_name = doc.add_paragraph()
    borrower_name_run = borrower_name.add_run('姓名：李广')
    borrower_name_run.font.size = Pt(12)
    set_chinese_font(borrower_name_run)

    borrower_id = doc.add_paragraph()
    borrower_id_run = borrower_id.add_run('身份证号：________')
    borrower_id_run.font.size = Pt(12)
    set_chinese_font(borrower_id_run)

    # 出借人信息
    lender_info = doc.add_paragraph()
    lender_info_run = lender_info.add_run('出借人信息：')
    lender_info_run.font.size = Pt(12)
    lender_info_run.font.bold = True
    set_chinese_font(lender_info_run)

    lender_name = doc.add_paragraph()
    lender_name_run = lender_name.add_run('姓名：张利顺')
    lender_name_run.font.size = Pt(12)
    set_chinese_font(lender_name_run)

    lender_id = doc.add_paragraph()
    lender_id_run = lender_id.add_run('身份证号：________')
    lender_id_run.font.size = Pt(12)
    set_chinese_font(lender_id_run)

    # 结语和签名
    doc.add_paragraph()  # 空行

    ending = doc.add_paragraph()
    ending_run = ending.add_run('立此为据，空口无凭。')
    ending_run.font.size = Pt(12)
    ending_run.font.bold = True
    set_chinese_font(ending_run)

    # 签名和日期
    doc.add_paragraph()  # 空行

    # 创建签名行表格
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False

    # 设置列宽
    table.columns[0].width = Inches(4)
    table.columns[1].width = Inches(2.5)

    # 填充内容
    signature_cell = table.cell(0, 0)
    signature_cell.text = '借款人（亲笔签名并捺印）：________'

    date_cell = table.cell(0, 1)
    date_cell.text = '签订日期：2024年____月____日'

    # 设置单元格字体
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(12)
                    set_chinese_font(run)

    # 保存文档
    doc.save(filename)
    print(f"✓ Word文档已生成: {filename}")
    print(f"✓ 文件位置: {os.path.abspath(filename)}")

def set_chinese_font(run):
    """设置中文字体"""
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

if __name__ == "__main__":
    print("=" * 50)
    print("借条Word文档生成器")
    print("=" * 50)

    # 检查是否安装了python-docx
    try:
        import docx
    except ImportError:
        print("❌ 错误: 未安装python-docx库")
        print("请运行: pip3 install python-docx")
        exit(1)

    # 生成Word文档
    create_iou_docx()
    print("\n✓ 完成！")
